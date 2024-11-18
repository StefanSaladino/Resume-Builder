from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from pymongo import MongoClient
from bson import ObjectId
from docx import Document
from dotenv import load_dotenv
import os
import tempfile
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load environment variables from the .env file
load_dotenv()

app = Flask(__name__)

# Enable CORS for specific origins
CORS(app, origins=[
    'http://localhost:4200',
    'https://resume-builder-3aba3.web.app',
    'https://resume-builder-backend-ahjg.onrender.com'
], supports_credentials=True)

# Fetch sensitive data from environment variables
mongodb_connection_string = os.getenv('CONNECTION_STRING_MONGODB')
mongodb_db = os.getenv('DB')
mongodb_users = os.getenv('COLLECTION')

# Connect to MongoDB
client = MongoClient(mongodb_connection_string)
db = client[mongodb_db]
users = db[mongodb_users]

@app.before_request
def handle_options():
    if request.method == 'OPTIONS':
        response = jsonify({"status": "CORS preflight handled"})
        response.status_code = 200
        response.headers.add("Access-Control-Allow-Origin", request.headers.get("Origin"))
        response.headers.add("Access-Control-Allow-Methods", "GET,POST,OPTIONS")
        response.headers.add("Access-Control-Allow-Headers", "Content-Type,Authorization")
        return response

@app.route('/generate-doc', methods=['POST'])
def generate_doc():
    try:
        # Get the user ID from the request payload
        user_id = request.json.get('userId')
        if not user_id:
            return jsonify({"error": "User ID is required"}), 400

        # Fetch user from the database
        user = users.find_one({'_id': ObjectId(user_id)})
        if not user or 'resume' not in user or 'generatedResume' not in user['resume']:
            return jsonify({"error": "No resume found for this user"}), 404

        generated_resume = user['resume']['generatedResume']

        # Create a new Word document
        doc = Document()

        # Overall font and style settings
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # HEADER - Add Name and Contact Information Once
        header = doc.add_paragraph()
        name = header.add_run(f"{user['resume']['basicInfo']['firstName']} {user['resume']['basicInfo']['lastName']}\n")
        name.font.size = Pt(16)
        name.font.bold = True
        contact = header.add_run(
            f"{user['resume']['basicInfo'].get('emailAddress', '')} | {user['resume']['basicInfo'].get('phone', '')}\n"
        )
        contact.font.size = Pt(10)
        contact.font.italic = True
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Reduce spacing after the header
        header_format = header.paragraph_format
        header_format.space_after = Pt(6)

        # SECTIONS: Format dynamically generated resume content
        sections = generated_resume.split('\n\n')  # Assuming '\n\n' separates sections
        for i, section in enumerate(sections):
            lines = section.split('\n')

            # Title of the section
            if len(lines) > 0:
                title = doc.add_paragraph(lines[0], style='Heading 1')
                title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # Reduce spacing after the title
                title_format = title.paragraph_format
                title_format.space_after = Pt(6)

                # Content of the section
                for line in lines[1:]:
                    if line.strip():
                        if line.startswith("-"):
                            content_paragraph = doc.add_paragraph(line.strip("- "), style='List Bullet')
                        else:
                            content_paragraph = doc.add_paragraph(line.strip())

                        # Reduce spacing after each paragraph
                        content_format = content_paragraph.paragraph_format
                        content_format.space_after = Pt(3)

            # Add spacing between sections, except after the last section
            if i < len(sections) - 1:
                doc.add_paragraph("")

        # FOOTER - Add References Information Once
        footer = doc.add_paragraph()
        footer.add_run("\nReferences available upon request").italic = True
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save the document to a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)

        # Return the document as a file download
        response = send_file(temp_file.name, as_attachment=True, download_name="resume.docx")
        response.headers.add("Access-Control-Allow-Origin", request.headers.get("Origin"))
        return response

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": f"An error occurred: {e}"}), 500


if __name__ == '__main__':
    app.run(debug=True)
